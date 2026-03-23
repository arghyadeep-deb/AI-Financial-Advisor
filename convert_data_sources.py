import os
import json
import pandas as pd
from pathlib import Path

DATA_DIR = Path("data_sources")
DATA_DIR.mkdir(exist_ok=True)


# ─── PDF Converter ────────────────────────────────────────────────────────────

def convert_pdf(filename: str, output_name: str):
    """Convert PDF inside data_sources/ to clean text."""
    pdf_path = DATA_DIR / filename

    if not pdf_path.exists():
        print(f"⚠️  Not found: {pdf_path} — skipping")
        return

    try:
        from pypdf import PdfReader

        reader = PdfReader(str(pdf_path))
        text   = ""

        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += f"\n--- Page {i+1} ---\n"
                text += page_text.strip()
                text += "\n"

        output_path = DATA_DIR / output_name
        output_path.write_text(text, encoding="utf-8")
        print(f"✅ PDF converted: {filename} → {output_name} ({len(text):,} chars, {len(reader.pages)} pages)")

    except ImportError:
        print("❌ pypdf not installed. Run: pip install pypdf")
    except Exception as e:
        print(f"❌ Failed to convert {filename}: {e}")


# ─── DOCX Converter ───────────────────────────────────────────────────────────

def convert_docx(filename: str, output_name: str):
    """Convert Word document inside data_sources/ to clean text."""
    docx_path = DATA_DIR / filename

    if not docx_path.exists():
        print(f"⚠️  Not found: {docx_path} — skipping")
        return

    try:
        import docx

        doc  = docx.Document(str(docx_path))
        text = ""

        for para in doc.paragraphs:
            if para.text.strip():
                # Detect headings
                if para.style.name.startswith("Heading"):
                    text += f"\n=== {para.text.strip()} ===\n"
                else:
                    text += para.text.strip() + "\n\n"

        # Extract tables too
        for table in doc.tables:
            text += "\n--- Table ---\n"
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    text += row_text + "\n"
            text += "\n"

        output_path = DATA_DIR / output_name
        output_path.write_text(text, encoding="utf-8")
        print(f"✅ DOCX converted: {filename} → {output_name} ({len(text):,} chars)")

    except ImportError:
        print("❌ python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        print(f"❌ Failed to convert {filename}: {e}")


# ─── XLSX Converter ───────────────────────────────────────────────────────────

def convert_xlsx(filename: str, output_name: str):
    """
    Convert Excel file inside data_sources/ to clean text.
    Formats each row as a readable paragraph — not raw table.
    """
    xlsx_path = DATA_DIR / filename

    if not xlsx_path.exists():
        print(f"⚠️  Not found: {xlsx_path} — skipping")
        return

    try:
        xl   = pd.ExcelFile(str(xlsx_path))
        text = f"Source: {filename}\n\n"

        for sheet_name in xl.sheet_names:
            df = pd.read_excel(str(xlsx_path), sheet_name=sheet_name)
            df = df.dropna(how="all")

            text += f"=== {sheet_name} ===\n\n"

            for _, row in df.iterrows():
                # Format each row as a readable sentence block
                row_text = ""
                for col in df.columns:
                    val = row.get(col, "")
                    if pd.notna(val) and str(val).strip():
                        row_text += f"{col}: {str(val).strip()}. "

                if row_text.strip():
                    text += row_text.strip() + "\n\n"

        output_path = DATA_DIR / output_name
        output_path.write_text(text, encoding="utf-8")
        rows_total = sum(
            len(pd.read_excel(str(xlsx_path), sheet_name=s))
            for s in xl.sheet_names
        )
        print(f"✅ XLSX converted: {filename} → {output_name} ({rows_total} rows, {len(text):,} chars)")

    except Exception as e:
        print(f"❌ Failed to convert {filename}: {e}")


# ─── Main ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("=" * 55)
    print("  Converting data_sources/ files to text")
    print("=" * 55)

    # ── All your files — already inside data_sources/ ────────────────────
    conversions = [
        {
            "input":  "Credit Card Project (1).pdf",
            "output": "credit_card_project.txt",
            "type":   "pdf"
        },
        {
            "input":  "CreditCardDetails (1).pdf",
            "output": "credit_card_details.txt",
            "type":   "pdf"
        },
        {
            "input":  "Investment.pdf",
            "output": "investment_guide.txt",
            "type":   "pdf"
        },
        {
            "input":  "financial_knowledge_base.docx",
            "output": "financial_knowledge_base.txt",
            "type":   "docx"
        },
        {
            "input":  "Indian Credit Card Report 2026 (1).xlsx",
            "output": "indian_credit_card_report_2026.txt",
            "type":   "xlsx"
        }
    ]

    for item in conversions:
        if item["type"] == "pdf":
            convert_pdf(item["input"], item["output"])
        elif item["type"] == "docx":
            convert_docx(item["input"], item["output"])
        elif item["type"] == "xlsx":
            convert_xlsx(item["input"], item["output"])

    print("=" * 55)
    print("✅ All conversions complete!")
    print()

    # ── Show what was created ─────────────────────────────────────────────
    txt_files = list(DATA_DIR.glob("*.txt"))
    if txt_files:
        print("Files created in data_sources/:")
        for f in sorted(txt_files):
            size = f.stat().st_size
            print(f"  📄 {f.name} ({size:,} bytes)")
    print()
    print("Now run: python rebuild_graph.py")
    print("=" * 55)